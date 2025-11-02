#!/usr/bin/env python3
"""
Export Agent - DOCX→PDF via LibreOffice
Replaces ReportLab with DOCX generation and LibreOffice conversion for better Unicode support.
"""
import sys
import json
import os
import subprocess
import time
import platform
import shlex
import uuid
import shutil
from pathlib import Path
from typing import Dict, Any, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    print("WARNING: python-docx not installed. Install with: pip install python-docx", file=sys.stderr)

AGENT_VERSION = "1.0.0"
AGENT_ID = "export_agent"

def export_docx(report: Dict[str, Any], out_dir: str = "out") -> str:
    """Generate DOCX file from report."""
    if Document is None:
        raise RuntimeError("python-docx not available. Install: pip install python-docx")
    
    os.makedirs(out_dir, exist_ok=True)
    docx_path = os.path.join(out_dir, "report.docx")
    
    doc = Document()
    
    # Set font to DejaVu Sans (or Calibri which has good Unicode support)
    # This ensures å/ä/ö render correctly
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'  # Fallback: Calibri has good Unicode support
        font._element.set(qn('w:eastAsia'), 'Calibri')  # For Asian chars
    except Exception:
        pass  # Font setting is optional
    
    # Title
    doc.add_heading("Relationsrapport", 0)
    
    # Summary
    s = report.get("summary", {})
    doc.add_heading("Sammanfattning", 1)
    doc.add_paragraph(f"Case ID: {s.get('case_id', 'unknown')}")
    doc.add_paragraph(f"Timestamp: {s.get('timestamp', 'N/A')}")
    
    overall_score = s.get("overall_score")
    if overall_score is not None:
        doc.add_paragraph(f"Övergripande poäng: {overall_score:.3f}")
    else:
        doc.add_paragraph("Övergripande poäng: N/A")
    
    doc.add_paragraph(f"Säkerhet: {s.get('safety_status', 'OK')}")
    focus = s.get("focus_areas", [])
    if focus:
        doc.add_paragraph(f"Fokusområden: {', '.join(focus)}")
    
    # Insights
    insights = report.get("insights", {})
    if insights:
        doc.add_heading("Insikter", 1)
        for key, value in insights.items():
            if value:
                doc.add_paragraph(f"{key}: {json.dumps(value, ensure_ascii=False)}")
    
    # Plan
    plan = report.get("plan", {})
    if plan:
        doc.add_heading("Interventionsplan", 1)
        title = plan.get("title", "Plan")
        doc.add_paragraph(title)
        interventions = plan.get("interventions", [])
        for iv in interventions:
            doc.add_heading(iv.get("title", "Intervention"), 2)
            for activity in iv.get("activities", []):
                doc.add_paragraph(f"• {activity}", style="List Bullet")
    
    # Recommendations
    recs = report.get("recommendations", [])
    if recs:
        doc.add_heading("Rekommendationer", 1)
        for rec in recs:
            doc.add_paragraph(rec, style="List Bullet")
    
    # Next steps
    next_steps = report.get("next_steps", [])
    if next_steps:
        doc.add_heading("Nästa steg", 1)
        for step in next_steps:
            doc.add_paragraph(step, style="List Bullet")
    
    doc.save(docx_path)
    os.utime(docx_path, None)  # bump mtime så LO ser färsk fil
    time.sleep(0.05)
    return docx_path

def _soffice_path():
    """Find LibreOffice soffice executable."""
    p = os.getenv("LIBREOFFICE_PATH")
    if p and os.path.exists(p):
        return p
    return "soffice.exe" if platform.system() == "Windows" else "soffice"


def _pdf_ok(path: str) -> bool:
    """Check if PDF file is valid and accessible."""
    try:
        if not os.path.exists(path): 
            return False
        if os.path.getsize(path) <= 1024:  # skydda mot tom/korrupt
            return False
        with open(path, "rb") as f:  # lås-koll på Windows
            f.read(4)
        return True
    except Exception:
        return False


def convert_to_pdf(docx_path: str, backend: str = "libreoffice", final_name: str = "report.pdf") -> Tuple[str, str]:
    """Convert DOCX to PDF using LibreOffice with retry and polling.
    Returns: (pdf_path, backend_used)
    """
    if backend == "docx-only":
        # Skip PDF generation, return DOCX path
        return docx_path, "docx-only"
    
    docx_path = os.path.abspath(docx_path)
    outdir = os.path.dirname(docx_path)
    soffice = _soffice_path()

    # unik temporär PDF (minskar lock-krockar på Windows)
    tmp_pdf = os.path.join(outdir, f"{uuid.uuid4().hex}.pdf")
    args = [
        soffice, "--headless", "--norestore", "--nodefault",
        "--nolockcheck", "--nofirststartwizard",
        "--convert-to", "pdf:writer_pdf_Export",  # explicit filter
        docx_path, "--outdir", outdir
    ]

    env = {**os.environ, "PYTHONIOENCODING": "utf-8"}
    
    # 1) Kör konvertering
    try:
        proc = subprocess.run(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE,
                              text=True, encoding="utf-8", errors="replace", env=env, timeout=30)
        
        if proc.returncode != 0:
            raise RuntimeError(f"LibreOffice konvertering misslyckades (rc={proc.returncode})\nSTDERR:\n{proc.stderr}")
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        # LibreOffice not found or timeout, fallback to docx-only
        return docx_path, "docx-only"

    # 2) Hitta genererad PDF (LibreOffice namnger enligt DOCX-basnamnet)
    base_pdf = docx_path.replace(".docx", ".pdf")
    cand = base_pdf if os.path.exists(base_pdf) else os.path.join(outdir, os.path.basename(base_pdf))

    # 3) Polling + retry mot fil-lås/timing
    deadline = time.time() + 15  # upp till 15s
    last_err = ""
    while time.time() < deadline:
        if os.path.exists(cand):
            # försök att "stabilisera" filen (Windows lås)
            time.sleep(0.3 if platform.system() == "Windows" else 0.1)
            if _pdf_ok(cand):
                # flytta till unik tmp och sedan till final
                try:
                    shutil.move(cand, tmp_pdf)
                    break
                except Exception:
                    # File might be locked, wait and retry
                    time.sleep(0.2)
                    continue
        time.sleep(0.3)

    if not _pdf_ok(tmp_pdf):
        # PDF conversion failed
        # If EXPORT_BACKEND=libreoffice is set, raise error instead of fallback
        if os.getenv("EXPORT_BACKEND", "").lower() == "libreoffice":
            raise RuntimeError(
                f"PDF conversion failed (backend=libreoffice required). "
                f"Expected PDF at {cand} but not found or invalid. "
                f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
            )
        # Fallback to docx-only if backend is not explicitly libreoffice
        if os.path.exists(tmp_pdf):
            try:
                os.remove(tmp_pdf)
            except Exception:
                pass
        return docx_path, "docx-only"

    final_pdf = os.path.join(outdir, final_name)
    # atomisk ersättning
    if os.path.exists(final_pdf):
        try:
            os.remove(final_pdf)
        except Exception:
            pass
    shutil.move(tmp_pdf, final_pdf)

    # extra paus på Windows för att undvika efter-lås
    if platform.system() == "Windows":
        time.sleep(0.2)

    return final_pdf, "libreoffice"

def run(payload: Dict[str, Any]) -> Dict[str, Any]:
    """Run export agent."""
    data = payload.get("data", {}) or {}
    meta = payload.get("meta", {}) or {}
    
    report = data.get("report", {})
    if not report:
        return {
            "ok": False,
            "error": "Missing report in data.report"
        }
    
    out_dir = meta.get("out_dir", "out")
    backend = os.getenv("EXPORT_BACKEND", meta.get("export_backend", "libreoffice")).lower()
    
    # Validate backend
    if backend not in ("libreoffice", "docx-only"):
        backend = "libreoffice"
    
    try:
        docx_path = export_docx(report, out_dir)
        
        # Ensure DOCX is not empty
        if os.path.exists(docx_path):
            docx_size = os.path.getsize(docx_path)
            if docx_size == 0:
                raise RuntimeError("Generated DOCX is empty")
        
        pdf_path, backend_used = convert_to_pdf(docx_path, backend, "report.pdf")
        
        pdf_exists = pdf_path.endswith(".pdf") and os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0
        
        return {
            "ok": True,
            "emits": {
                "docx_path": docx_path,
                "pdf_path": pdf_path if pdf_exists else None,
                "backend_used": backend_used
            },
            "checks": {
                "CHK-DOCX-EXISTS": {
                    "pass": os.path.exists(docx_path),
                    "path": docx_path,
                    "size_bytes": os.path.getsize(docx_path) if os.path.exists(docx_path) else 0
                },
                "CHK-PDF-EXISTS": {
                    "pass": pdf_exists,
                    "path": pdf_path if pdf_exists else None,
                    "size_bytes": os.path.getsize(pdf_path) if pdf_exists else 0
                }
            }
        }
    except Exception as e:
        return {
            "ok": False,
            "error": str(e)
        }

def main():
    t0 = time.time()
    try:
        payload_str = sys.stdin.read() or "{}"
        payload = json.loads(payload_str)
        result = run(payload)
        result["version"] = f"{AGENT_ID}@{AGENT_VERSION}"
        result["latency_ms"] = int((time.time() - t0) * 1000)
        result["cost"] = {"usd": 0.001}
        print(json.dumps(result, ensure_ascii=False))
    except Exception as e:
        print(json.dumps({"ok": False, "error": str(e)}, ensure_ascii=False))
        sys.exit(1)

if __name__ == "__main__":
    main()

